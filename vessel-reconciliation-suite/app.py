from __future__ import annotations

import io
import re
import json
import math
import hashlib
import zipfile
from dataclasses import dataclass, asdict
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import pandas as pd
import numpy as np

try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


# =========================
# Exceptions
# =========================

class ParserError(Exception):
    pass


class SchemaError(ParserError):
    pass


class ReconciliationError(ParserError):
    pass


# =========================
# Dataclasses
# =========================

@dataclass(frozen=True)
class SourceFingerprint:
    file_name: str
    sha256: str
    byte_size: int


@dataclass
class PMSItemRecord:
    vessel_name: str
    code_no: str
    item_name: str
    job: str
    interval_raw: str
    interval_hrs_1: Optional[float]
    interval_hrs_2: Optional[float]
    date_last_inspection: Optional[pd.Timestamp]
    op_hours_end_last_year: Optional[float]
    current_operating_hours: Optional[float]
    estimated_next_inspection: Optional[pd.Timestamp]
    jan: Optional[float]
    feb: Optional[float]
    mar: Optional[float]
    apr: Optional[float]
    may: Optional[float]
    jun: Optional[float]
    jul: Optional[float]
    aug: Optional[float]
    sep: Optional[float]
    oct: Optional[float]
    nov: Optional[float]
    dec: Optional[float]
    equipment_group: str
    section_name: str
    source_sheet: str
    source_row_excel: int


@dataclass
class RunningHoursWordRecord:
    vessel_name: str
    section: str
    cyl_or_unit: Optional[str]
    description: str
    periodicity_raw: str
    date_of_last_oh: Optional[pd.Timestamp]
    running_hours_since_last_oh: Optional[float]
    total_running_hours: Optional[float]
    this_month_hours: Optional[float]
    raw_excerpt: str


@dataclass
class LedgerRecord:
    ledger_id: str
    vessel_name: str
    code_no: str
    pms_item_name: str
    pms_section: str
    pms_group: str
    pms_date_last_inspection: Optional[pd.Timestamp]
    pms_current_operating_hours: Optional[float]
    pms_estimated_next_inspection: Optional[pd.Timestamp]
    pms_mar_hours: Optional[float]
    word_section: Optional[str]
    word_unit: Optional[str]
    word_description: Optional[str]
    word_date_last_oh: Optional[pd.Timestamp]
    word_running_hours_since_last_oh: Optional[float]
    word_total_running_hours: Optional[float]
    word_this_month_hours: Optional[float]
    match_score: float
    reconciliation_status: str
    reconciliation_notes: str
    pms_source_hash: str
    word_source_hash: Optional[str]
    created_at_utc: str


# =========================
# Constants
# =========================

MONTH_COLS = ["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"]

PMS_REQUIRED_CANONICAL = [
    "code_no",
    "items",
    "job",
    "interval_1",
    "interval_2",
    "date_last_inspection",
    "op_hours_end_last_year",
    "current_operating_hours",
    "estimated_date_next_inspection",
]

PMS_HEADER_SYNONYMS = {
    "code_no": {
        "code no", "code no.", "code", "code number"
    },
    "items": {
        "items", "item", "description"
    },
    "job": {
        "job"
    },
    "interval_1": {
        "interval"
    },
    "interval_2": {
        ""  # second interval column is often blank-headed
    },
    "date_last_inspection": {
        "date of last inspection", "date of last oh", "date last inspection"
    },
    "op_hours_end_last_year": {
        "operating hours at the end of last year",
        "operating hrs at the end of last year"
    },
    "current_operating_hours": {
        "current operating hours", "current op. hours", "current op hours"
    },
    "estimated_date_next_inspection": {
        "estimated date of next inspection", "estimated date next inspection"
    },
    "jan": {"jan", "jan."},
    "feb": {"feb", "feb."},
    "mar": {"mar", "mar."},
    "apr": {"apr", "apr."},
    "may": {"may", "may."},
    "jun": {"jun", "jun."},
    "jul": {"jul", "jul."},
    "aug": {"aug", "aug."},
    "sep": {"sep", "sep."},
    "oct": {"oct", "oct."},
    "nov": {"nov", "nov."},
    "dec": {"dec", "dec."},
}

WORD_SECTION_ALIASES = {
    "main engine": "MAIN ENGINE",
    "turbocharger": "TURBOCHARGER",
    "auxiliary boiler": "AUXILIARY BOILER",
    "exh gas boiler": "EXHAUST GAS BOILER",
    "main air compressors": "MAIN AIR COMPRESSORS",
    "aux. engine": "AUX ENGINE",
    "aux engine": "AUX ENGINE",
    "dg": "DIESEL GENERATOR",
    "diesel generator": "DIESEL GENERATOR",
}

ITEM_NORMALIZATION_MAP = {
    "cylinder cover": "CYLINDER COVER",
    "cylinder cover cooling jacket": "CYLINDER COVER COOLING JACKET",
    "cyl. liner": "CYLINDER LINER",
    "cyl liner": "CYLINDER LINER",
    "cylinder liner": "CYLINDER LINER",
    "piston assy": "PISTON ASSEMBLY",
    "piston assembly": "PISTON ASSEMBLY",
    "stuffing box": "STUFFING BOX",
    "piston crown": "PISTON CROWN",
    "exhaust valve": "EXHAUST VALVE",
    "exaust valve": "EXHAUST VALVE",
    "starting valve": "STARTING VALVE",
    "safety valve": "SAFETY VALVE",
    "fuel valve": "FUEL VALVES",
    "fuel valves": "FUEL VALVES",
    "fuel pump": "FUEL PUMP",
    "plunger and barrel renewal": "PLUNGER AND BARREL RENEWAL",
    "fuel pump suction valve": "FUEL PUMP SUCTION VALVE",
    "fuel pump puncture valve": "FUEL PUMP PUNCTURE VALVE",
    "crosshead pin bearing": "CROSSHEAD BEARINGS",
    "crosshead bearings": "CROSSHEAD BEARINGS",
    "crankpin bearing": "BOTTOM END BEARINGS",
    "bottom end bearings": "BOTTOM END BEARINGS",
    "main bearings": "MAIN BEARINGS",
    "main bearing": "MAIN BEARINGS",
    "adjust valve head clearance": "ADJUST VALVE HEAD CLEARANCE",
    "turbocharger general inspection": "TURBOCHARGER",
    "air cooler cleaning": "AIR COOLER CLEANING",
}


# =========================
# Utility
# =========================

def file_fingerprint(file_name: str, file_bytes: bytes) -> SourceFingerprint:
    return SourceFingerprint(
        file_name=file_name,
        sha256=hashlib.sha256(file_bytes).hexdigest(),
        byte_size=len(file_bytes)
    )


def normalize_text(x) -> str:
    if x is None:
        return ""
    s = str(x)
    s = s.replace("\xa0", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def normalize_key(x) -> str:
    s = normalize_text(x).lower()
    s = s.replace("_", " ")
    s = re.sub(r"[^\w\s./-]", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def normalize_item_name(x: str) -> str:
    s = normalize_key(x)
    s = s.replace("insp. through scav ports", "")
    s = s.replace("through scav ports", "")
    s = s.replace("p. rings", "")
    s = s.replace("cooling jacket", "cooling jacket")
    s = re.sub(r"\bno\.?\s*\d+\b", "", s)
    s = re.sub(r"\bcyl\.?\b", "cylinder", s)
    s = re.sub(r"\bassy\b", "assembly", s)
    s = re.sub(r"\s+", " ", s).strip()
    for k, v in ITEM_NORMALIZATION_MAP.items():
        if k in s:
            return v
    return s.upper()


def to_float(x) -> Optional[float]:
    if x is None:
        return None
    if isinstance(x, float) and math.isnan(x):
        return None
    s = normalize_text(x)
    if s == "":
        return None
    s = s.replace(",", "")
    s = s.replace("HRS", "").replace("Hours", "").replace("hours", "")
    s = s.strip()
    if s in {"-", "--", "NA", "N/A"}:
        return None
    try:
        return float(s)
    except Exception:
        return None


def to_timestamp(x) -> Optional[pd.Timestamp]:
    if x is None:
        return None
    if isinstance(x, pd.Timestamp):
        return x if not pd.isna(x) else None
    s = normalize_text(x)
    if s == "":
        return None

    s = s.replace(" -", "-").replace("- ", "-")
    s = re.sub(r"\bSEPT\b", "SEP", s, flags=re.I)
    s = re.sub(r"\bJULY\b", "JUL", s, flags=re.I)
    s = re.sub(r"\bJUNE\b", "JUN", s, flags=re.I)
    s = re.sub(r"\bMARCH\b", "MAR", s, flags=re.I)
    s = re.sub(r"\s+", " ", s).strip()

    for dayfirst in (True, False):
        try:
            ts = pd.to_datetime(s, errors="raise", dayfirst=dayfirst)
            if pd.isna(ts):
                continue
            return ts
        except Exception:
            pass

    return None


def safe_str(x) -> str:
    return "" if x is None or (isinstance(x, float) and math.isnan(x)) else str(x)


def is_code_like(s: str) -> bool:
    s = normalize_text(s).upper()
    return bool(re.match(r"^[A-Z]{2,4}-\d{2}(?:-\d{2}(?:\.\d+)?)?$", s))


def vessel_from_sheet(df: pd.DataFrame) -> str:
    sample = " ".join(
        normalize_text(v)
        for v in df.iloc[:12, :8].fillna("").astype(str).values.ravel().tolist()
        if normalize_text(v)
    )
    m = re.search(r"MV\s+([A-Z0-9 .\-]+)", sample, flags=re.I)
    if not m:
        raise SchemaError("Unable to locate vessel name in workbook top area.")
    return normalize_text("MV " + m.group(1))


def find_sheet_name(xls: pd.ExcelFile, preferred: str) -> str:
    names = list(xls.sheet_names)
    exact = [n for n in names if normalize_key(n) == normalize_key(preferred)]
    if exact:
        return exact[0]
    partial = [n for n in names if normalize_key(preferred) in normalize_key(n)]
    if partial:
        return partial[0]
    raise SchemaError(f"Required sheet '{preferred}' not found. Available sheets: {names}")


def row_text(row: pd.Series) -> str:
    return " | ".join(normalize_text(v) for v in row.tolist() if normalize_text(v))


# =========================
# DAILY OPERATING HOURS
# =========================

def parse_daily_operating_hours(file_bytes: bytes) -> Tuple[str, pd.DataFrame, Dict[str, float]]:
    xls = pd.ExcelFile(io.BytesIO(file_bytes), engine="openpyxl")
    sheet = find_sheet_name(xls, "DAILY OPERATING HOURS")
    raw = pd.read_excel(xls, sheet_name=sheet, header=None)

    vessel_name = vessel_from_sheet(raw)

    header_row = None
    machine_row = None
    for i in range(min(len(raw) - 1, 40)):
        line1 = row_text(raw.iloc[i])
        line2 = row_text(raw.iloc[i + 1]) if i + 1 < len(raw) else ""
        if "DATE" in line1.upper() and "OPERATING HOURS" in line1.upper():
            header_row = i
            machine_row = i + 1
            break

    if header_row is None or machine_row is None:
        raise SchemaError("Could not anchor DAILY OPERATING HOURS header block.")

    date_col = None
    machine_map: Dict[int, str] = {}
    mach_line = [normalize_text(v) for v in raw.iloc[machine_row].tolist()]

    for j, v in enumerate(mach_line):
        vu = v.upper()
        if vu == "DATE":
            date_col = j
        if "MAIN ENGINE" in vu:
            machine_map[j] = "ME"
        elif "DIESEL GENERATOR NO.1" in vu or "DIESEL GENERATOR NO. 1" in vu:
            machine_map[j] = "DG1"
        elif "DIESEL GENERATOR NO.2" in vu or "DIESEL GENERATOR NO. 2" in vu:
            machine_map[j] = "DG2"
        elif "DIESEL GENERATOR NO.3" in vu or "DIESEL GENERATOR NO. 3" in vu:
            machine_map[j] = "DG3"
        elif "DIESEL GENERATOR NO.4" in vu or "DIESEL GENERATOR NO. 4" in vu:
            machine_map[j] = "DG4"

    if date_col is None:
        date_col = 0

    if not machine_map:
        raise SchemaError("Could not map machines in DAILY OPERATING HOURS sheet.")

    records = []
    monthly_totals: Dict[str, float] = {}

    start = machine_row + 1
    for i in range(start, len(raw)):
        first = normalize_text(raw.iat[i, date_col])
        if first == "":
            continue

        if first.lower().startswith("total monthly oper"):
            for col_idx, machine in machine_map.items():
                val = to_float(raw.iat[i, col_idx + 1] if col_idx + 1 < raw.shape[1] else None)
                if val is not None:
                    month_key = f"{machine}_row_{i}"
                    monthly_totals[month_key] = val
            continue

        dt = to_timestamp(first)
        if dt is None:
            continue

        for col_idx, machine in sorted(machine_map.items()):
            val = to_float(raw.iat[i, col_idx + 1] if col_idx + 1 < raw.shape[1] else None)
            if val is not None:
                records.append({
                    "date": dt.normalize(),
                    "machine": machine,
                    "hours": val,
                    "source_row_excel": i + 1,
                    "source_sheet": sheet,
                })

    df = pd.DataFrame(records)
    if df.empty:
        raise SchemaError("No daily operating hours records extracted.")

    df["month"] = df["date"].dt.month
    agg = (
        df.groupby(["machine", "month"], as_index=False)["hours"]
        .sum()
        .rename(columns={"hours": "monthly_hours"})
    )

    summary = {
        f"{r.machine}_M{int(r.month):02d}": float(r.monthly_hours)
        for r in agg.itertuples(index=False)
    }

    return vessel_name, df, summary


# =========================
# PMS SHEET
# =========================

def find_pms_header_row(raw: pd.DataFrame) -> int:
    for i in range(min(len(raw), 80)):
        vals = [normalize_key(v) for v in raw.iloc[i].tolist()]
        joined = " | ".join(vals)
        if (
            "code no" in joined
            and "items" in joined
            and "date of last inspection" in joined
            and "current operating hours" in joined
        ):
            return i
    raise SchemaError("PMS header row not found.")


def canonicalize_pms_headers(raw: pd.DataFrame, header_row: int) -> Dict[str, int]:
    vals = [normalize_key(v) for v in raw.iloc[header_row].tolist()]
    mapping: Dict[str, int] = {}

    for idx, val in enumerate(vals):
        for canon, synonyms in PMS_HEADER_SYNONYMS.items():
            if val in synonyms:
                if canon not in mapping:
                    mapping[canon] = idx

    # Special handling for duplicate interval columns next to JOB
    if "interval_1" not in mapping:
        raise SchemaError("PMS column 'INTERVAL' not found.")
    first_interval_idx = mapping["interval_1"]
    mapping["interval_2"] = first_interval_idx + 1

    missing = [c for c in PMS_REQUIRED_CANONICAL if c not in mapping]
    if missing:
        raise SchemaError(f"Missing canonical PMS columns: {missing}")

    for m in MONTH_COLS:
        if m not in mapping:
            raise SchemaError(f"Missing monthly PMS column: {m}")

    return mapping


def classify_pms_row(code_no: str, items: str) -> str:
    code = normalize_text(code_no).upper()
    item = normalize_text(items).upper()

    if not code and not item:
        return "blank"
    if is_code_like(code):
        return "item"
    if code.endswith("-00") or code.endswith("-00-00"):
        return "section"
    if "TOTAL WORKING HOURS" in code or "CURRENT YEAR UP-TO-DATE" in code:
        return "summary"
    if item == "" and code != "":
        return "section"
    return "other"


def parse_pms_sheet(file_bytes: bytes) -> Tuple[str, pd.DataFrame]:
    xls = pd.ExcelFile(io.BytesIO(file_bytes), engine="openpyxl")
    sheet = find_sheet_name(xls, "PMS")
    raw = pd.read_excel(xls, sheet_name=sheet, header=None)

    vessel_name = vessel_from_sheet(raw)
    header_row = find_pms_header_row(raw)
    colmap = canonicalize_pms_headers(raw, header_row)

    current_group = ""
    current_section = ""
    records: List[PMSItemRecord] = []

    for i in range(header_row + 1, len(raw)):
        code_no = normalize_text(raw.iat[i, colmap["code_no"]])
        items = normalize_text(raw.iat[i, colmap["items"]])

        row_type = classify_pms_row(code_no, items)

        if row_type == "blank":
            continue

        if row_type == "section":
            if code_no:
                current_section = code_no
            if items:
                current_group = items
            elif code_no and not items:
                current_group = code_no
            continue

        if row_type != "item":
            continue

        code_u = code_no.upper()
        if not is_code_like(code_u):
            continue

        rec = PMSItemRecord(
            vessel_name=vessel_name,
            code_no=code_u,
            item_name=items,
            job=normalize_text(raw.iat[i, colmap["job"]]),
            interval_raw=" | ".join(
                [normalize_text(raw.iat[i, colmap["interval_1"]]), normalize_text(raw.iat[i, colmap["interval_2"]])]
            ).strip(" |"),
            interval_hrs_1=to_float(raw.iat[i, colmap["interval_1"]]),
            interval_hrs_2=to_float(raw.iat[i, colmap["interval_2"]]),
            date_last_inspection=to_timestamp(raw.iat[i, colmap["date_last_inspection"]]),
            op_hours_end_last_year=to_float(raw.iat[i, colmap["op_hours_end_last_year"]]),
            current_operating_hours=to_float(raw.iat[i, colmap["current_operating_hours"]]),
            estimated_next_inspection=to_timestamp(raw.iat[i, colmap["estimated_date_next_inspection"]]),
            jan=to_float(raw.iat[i, colmap["jan"]]),
            feb=to_float(raw.iat[i, colmap["feb"]]),
            mar=to_float(raw.iat[i, colmap["mar"]]),
            apr=to_float(raw.iat[i, colmap["apr"]]),
            may=to_float(raw.iat[i, colmap["may"]]),
            jun=to_float(raw.iat[i, colmap["jun"]]),
            jul=to_float(raw.iat[i, colmap["jul"]]),
            aug=to_float(raw.iat[i, colmap["aug"]]),
            sep=to_float(raw.iat[i, colmap["sep"]]),
            oct=to_float(raw.iat[i, colmap["oct"]]),
            nov=to_float(raw.iat[i, colmap["nov"]]),
            dec=to_float(raw.iat[i, colmap["dec"]]),
            equipment_group=current_group,
            section_name=current_section,
            source_sheet=sheet,
            source_row_excel=i + 1,
        )
        records.append(rec)

    if not records:
        raise SchemaError("No PMS item rows extracted.")

    df = pd.DataFrame([asdict(r) for r in records])

    critical_nulls = df["code_no"].isna().sum() + df["item_name"].isna().sum()
    if critical_nulls > 0:
        raise SchemaError("Null critical values detected after PMS extraction.")

    return vessel_name, df


# =========================
# WORD / DOC PARSER
# =========================

def extract_doc_text(file_bytes: bytes, file_name: str) -> str:
    lower = file_name.lower()

    if lower.endswith(".docx"):
        if not HAS_DOCX:
            raise ParserError("python-docx is required to parse .docx files.")
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if normalize_text(p.text)]
        for tbl in doc.tables:
            for row in tbl.rows:
                cell_texts = [normalize_text(c.text) for c in row.cells]
                if any(cell_texts):
                    parts.append(" | ".join(cell_texts))
        return "\n".join(parts)

    if lower.endswith(".doc"):
        # Many .doc uploads in these workflows are actually text-extracted artifacts.
        try:
            txt = file_bytes.decode("utf-8", errors="ignore")
            if len(normalize_text(txt)) > 50:
                return txt
        except Exception:
            pass
        try:
            txt = file_bytes.decode("latin-1", errors="ignore")
            if len(normalize_text(txt)) > 50:
                return txt
        except Exception:
            pass
        raise ParserError("Binary .doc parsing is not supported in this core. Convert to .docx or text first.")

    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return file_bytes.decode("latin-1", errors="ignore")


def split_word_sections(text: str) -> List[str]:
    t = normalize_text(text)
    t = t.replace("TITLE Vessels Name", "\nTITLE Vessels Name")
    t = t.replace("TABLE Note 1", "\nTABLE Note 1")
    t = t.replace("AUXILIARY BOILER", "\nAUXILIARY BOILER")
    t = t.replace("TURBOCHARGER", "\nTURBOCHARGER")
    t = t.replace("DESCRIPTIONPERIODICITYDG No1DG No2DG No3", "\nDESCRIPTIONPERIODICITYDG No1DG No2DG No3")
    chunks = [c.strip() for c in re.split(r"\n+", t) if normalize_text(c)]
    return chunks


def parse_word_running_hours(file_bytes: bytes, file_name: str) -> Tuple[str, pd.DataFrame]:
    text = extract_doc_text(file_bytes, file_name)
    if "Running Hours Monthly Report" not in text and "RUNNING HOURS SINCE LAST OH" not in text:
        raise SchemaError("Input document does not look like the running-hours monthly report.")

    vessel_match = re.search(r"Vessels Name\s+(MV\s+[A-Z0-9 .\-]+)", text, flags=re.I)
    if not vessel_match:
        vessel_match = re.search(r"(MV\s+MINOAN\s+FALCON)", text, flags=re.I)
    if not vessel_match:
        raise SchemaError("Could not find vessel name in Word report.")
    vessel_name = normalize_text(vessel_match.group(1))

    me_total = None
    me_month = None
    m_total = re.search(r"MAIN ENGINE Type\s+Total Running Hours\s+([\d,]+)\s*Hours\s+This Month\s+([\d,]+)\s*HRS", text, flags=re.I)
    if m_total:
        me_total = to_float(m_total.group(1))
        me_month = to_float(m_total.group(2))

    records: List[RunningHoursWordRecord] = []

    # Main engine cylinder lines
    if "CYL. No.1" in text and "CYL. No.2" in text:
        for desc in [
            "CYLINDER COVER", "PISTON ASSEMBLY", "STUFFING BOX", "PISTON CROWN",
            "CYLINDER LINER", "EXAUST VALVE", "STARTING VALVE", "SAFETY VALVE",
            "FUEL VALVES", "FUEL PUMP", "PLUNGER AND BARREL RENEWAL",
            "FUEL PUMP SUCTION VALVE", "FUEL PUMP PUNCTURE VALVE",
            "CROSSHEAD BEARINGS", "BOTTOM END BEARINGS", "MAIN BEARINGS"
        ]:
            pattern = re.compile(
                rf"{re.escape(desc)}\s+(.+?)(?=(CYLINDER COVER|PISTON ASSEMBLY|STUFFING BOX|PISTON CROWN|CYLINDER LINER|EXAUST VALVE|STARTING VALVE|SAFETY VALVE|FUEL VALVES|FUEL PUMP|PLUNGER AND BARREL RENEWAL|FUEL PUMP SUCTION VALVE|FUEL PUMP PUNCTURE VALVE|CROSSHEAD BEARINGS|BOTTOM END BEARINGS|MAIN BEARINGS|Note 1))",
                flags=re.I | re.S
            )
            m = pattern.search(text)
            if not m:
                continue

            excerpt = normalize_text(m.group(1))
            dates = re.findall(r"\b\d{1,2}[-/][A-Z]{3,9}[-/]\d{2,4}\b|\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b", excerpt, flags=re.I)
            nums = re.findall(r"(?<![A-Z])\d[\d,]*", excerpt)

            dt = to_timestamp(dates[0]) if dates else None
            hrs = None
            cleaned_nums = [to_float(x) for x in nums if to_float(x) is not None]
            if cleaned_nums:
                hrs = cleaned_nums[-1]

            records.append(RunningHoursWordRecord(
                vessel_name=vessel_name,
                section="MAIN ENGINE",
                cyl_or_unit=None,
                description=desc,
                periodicity_raw="",
                date_of_last_oh=dt,
                running_hours_since_last_oh=hrs,
                total_running_hours=me_total,
                this_month_hours=me_month,
                raw_excerpt=excerpt[:500]
            ))

    # DG block
    dg_months = re.search(
        r"Total Hours\s*([0-9,]+).*?Hours This Month\s*([0-9,]+).*?Hours This Month.*?DESCRIPTIONPERIODICITYDG No1DG No2DG No3(.*?)(TABLE 1st Copy|TITLE Vessels Name)",
        text,
        flags=re.I | re.S
    )
    if dg_months:
        dg_total = to_float(dg_months.group(1))
        dg_this_month = to_float(dg_months.group(2))
        dg_body = dg_months.group(3)

        line_candidates = [
            "Cylinder Head", "Piston", "Connecting Rod", "Cylinder Liners", "Fuel Valves",
            "Fuel Pumps", "Crank Pin Bearing", "Main Bearing", "Adjust Valve Head Clearance",
            "Turbocharger", "Air Cooler", "L.O. Cooler Clean", "Cooling Water Pump",
            "F.W. Cooler Clean", "Cool Water Thermostat Valve", "L.O. Renewal",
            "Alternator Cleaning", "L.O. Thermostat Valve", "Thrust Bearing"
        ]

        for item in line_candidates:
            pat = re.compile(rf"{re.escape(item)}\s+(.*?)(?=(?:{'|'.join(map(re.escape, line_candidates))}|TABLE 1st Copy|$))", flags=re.I | re.S)
            m = pat.search(dg_body)
            if not m:
                continue
            excerpt = normalize_text(m.group(1))
            dates = re.findall(r"\b\d{1,2}[A-Z]{3}\d{2}\b|\b\d{1,2}[A-Z]{3}\d{4}\b|\b\d{1,2}-[A-Z]{3}-\d{2,4}\b|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", excerpt, flags=re.I)
            nums = [to_float(x) for x in re.findall(r"\d[\d,]*", excerpt) if to_float(x) is not None]
            hrs = nums[-1] if nums else None
            dt = to_timestamp(dates[0]) if dates else None

            records.append(RunningHoursWordRecord(
                vessel_name=vessel_name,
                section="DIESEL GENERATOR",
                cyl_or_unit=None,
                description=item.upper(),
                periodicity_raw="",
                date_of_last_oh=dt,
                running_hours_since_last_oh=hrs,
                total_running_hours=dg_total,
                this_month_hours=dg_this_month,
                raw_excerpt=excerpt[:500]
            ))

    df = pd.DataFrame([asdict(r) for r in records])
    if df.empty:
        raise SchemaError("No running-hours records extracted from Word report.")

    return vessel_name, df

# =========================
# Matching + Reconciliation
# =========================

def infer_word_target_from_pms(code_no: str, item_name: str, equipment_group: str) -> Tuple[str, Optional[str]]:
    code = normalize_text(code_no).upper()
    item = normalize_item_name(item_name)
    grp = normalize_text(equipment_group).upper()

    if code.startswith("ME-"):
        if re.match(r"^ME-\d{2}-", code):
            cyl = re.search(r"ME-(\d{2})-", code)
            unit = f"CYL. NO.{int(cyl.group(1))}" if cyl else None
            return "MAIN ENGINE", unit

    if code.startswith("DG-"):
        dg = re.search(r"DG-(\d{2})-", code)
        unit = f"DG NO{int(dg.group(1))}" if dg else None
        return "DIESEL GENERATOR", unit

    if "TURBOCHARGER" in grp or "TURBOCHARGER" in item:
        return "TURBOCHARGER", None

    if "AIR COMPRESSOR" in grp or "COMPRESSOR" in grp:
        return "MAIN AIR COMPRESSORS", None

    return "", None


def compute_match_score(pms_row: pd.Series, word_row: pd.Series) -> float:
    score = 0.0

    p_item = normalize_item_name(pms_row["item_name"])
    w_item = normalize_item_name(word_row["description"])

    p_section, p_unit = infer_word_target_from_pms(
        pms_row["code_no"],
        pms_row["item_name"],
        pms_row.get("equipment_group", "")
    )
    w_section = normalize_text(word_row["section"]).upper()

    if p_section and p_section == w_section:
        score += 40.0

    if p_item == w_item:
        score += 40.0
    elif p_item in w_item or w_item in p_item:
        score += 25.0

    p_dt = pms_row.get("date_last_inspection")
    w_dt = word_row.get("date_of_last_oh")
    if pd.notna(p_dt) and pd.notna(w_dt):
        day_gap = abs((pd.Timestamp(p_dt).normalize() - pd.Timestamp(w_dt).normalize()).days)
        if day_gap == 0:
            score += 15.0
        elif day_gap <= 3:
            score += 10.0
        elif day_gap <= 10:
            score += 5.0

    p_hrs = pms_row.get("current_operating_hours")
    w_hrs = word_row.get("running_hours_since_last_oh")
    if pd.notna(p_hrs) and pd.notna(w_hrs):
        diff = abs(float(p_hrs) - float(w_hrs))
        if diff == 0:
            score += 15.0
        elif diff <= 5:
            score += 10.0
        elif diff <= 25:
            score += 5.0

    return min(score, 100.0)


def best_word_match_for_pms(pms_row: pd.Series, word_df: pd.DataFrame) -> Tuple[Optional[pd.Series], float]:
    candidates = word_df.copy()

    target_section, _ = infer_word_target_from_pms(
        pms_row["code_no"],
        pms_row["item_name"],
        pms_row.get("equipment_group", "")
    )

    if target_section:
        sec = candidates["section"].fillna("").str.upper() == target_section
        if sec.any():
            candidates = candidates[sec].copy()

    if candidates.empty:
        return None, 0.0

    scored = []
    for idx, row in candidates.iterrows():
        scored.append((idx, compute_match_score(pms_row, row)))

    scored.sort(key=lambda x: x[1], reverse=True)
    best_idx, best_score = scored[0]
    if best_score < 45.0:
        return None, best_score

    return candidates.loc[best_idx], best_score


def reconcile_status(
    pms_current_hours: Optional[float],
    word_running_hours: Optional[float],
    pms_mar: Optional[float],
    word_month: Optional[float],
    match_score: float,
) -> Tuple[str, str]:
    notes = []

    if match_score < 45:
        return "UNMATCHED", "No sufficiently strong deterministic match in Word report."

    status = "MATCHED"

    if pms_current_hours is not None and word_running_hours is not None:
        diff = abs(float(pms_current_hours) - float(word_running_hours))
        notes.append(f"Current-vs-word OH diff={diff:.1f}")
        if diff > 25:
            status = "REVIEW"

    if pms_mar is not None and word_month is not None:
        mdiff = abs(float(pms_mar) - float(word_month))
        notes.append(f"March-hours diff={mdiff:.1f}")
        if mdiff > 25 and status == "MATCHED":
            status = "REVIEW"

    if not notes:
        notes.append("Matched by section/item only.")

    return status, "; ".join(notes)


# =========================
# Immutable ledger build
# =========================

def make_ledger_id(pms_source_hash: str, code_no: str, source_row_excel: int) -> str:
    raw = f"{pms_source_hash}|{code_no}|{source_row_excel}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:24]


def build_immutable_ledger(
    pms_df: pd.DataFrame,
    word_df: Optional[pd.DataFrame],
    pms_fingerprint: SourceFingerprint,
    word_fingerprint: Optional[SourceFingerprint],
) -> pd.DataFrame:
    ledger_rows: List[LedgerRecord] = []

    for _, pms_row in pms_df.iterrows():
        matched_word = None
        score = 0.0

        if word_df is not None and not word_df.empty:
            matched_word, score = best_word_match_for_pms(pms_row, word_df)

        if matched_word is not None:
            status, notes = reconcile_status(
                pms_current_hours=pms_row.get("current_operating_hours"),
                word_running_hours=matched_word.get("running_hours_since_last_oh"),
                pms_mar=pms_row.get("mar"),
                word_month=matched_word.get("this_month_hours"),
                match_score=score,
            )
        else:
            status, notes = ("UNMATCHED", "No Word counterpart found or score below threshold.")

        row = LedgerRecord(
            ledger_id=make_ledger_id(pms_fingerprint.sha256, pms_row["code_no"], int(pms_row["source_row_excel"])),
            vessel_name=pms_row["vessel_name"],
            code_no=pms_row["code_no"],
            pms_item_name=pms_row["item_name"],
            pms_section=pms_row.get("section_name", ""),
            pms_group=pms_row.get("equipment_group", ""),
            pms_date_last_inspection=pms_row.get("date_last_inspection"),
            pms_current_operating_hours=pms_row.get("current_operating_hours"),
            pms_estimated_next_inspection=pms_row.get("estimated_next_inspection"),
            pms_mar_hours=pms_row.get("mar"),
            word_section=matched_word.get("section") if matched_word is not None else None,
            word_unit=matched_word.get("cyl_or_unit") if matched_word is not None else None,
            word_description=matched_word.get("description") if matched_word is not None else None,
            word_date_last_oh=matched_word.get("date_of_last_oh") if matched_word is not None else None,
            word_running_hours_since_last_oh=matched_word.get("running_hours_since_last_oh") if matched_word is not None else None,
            word_total_running_hours=matched_word.get("total_running_hours") if matched_word is not None else None,
            word_this_month_hours=matched_word.get("this_month_hours") if matched_word is not None else None,
            match_score=round(score, 2),
            reconciliation_status=status,
            reconciliation_notes=notes,
            pms_source_hash=pms_fingerprint.sha256,
            word_source_hash=word_fingerprint.sha256 if word_fingerprint else None,
            created_at_utc=datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
        )
        ledger_rows.append(row)

    ledger_df = pd.DataFrame([asdict(x) for x in ledger_rows])

    if ledger_df["ledger_id"].duplicated().any():
        raise ReconciliationError("Ledger ID collision detected. Immutable key generation failed.")

    return ledger_df


# =========================
# Validation
# =========================

def validate_pms_vs_daily_hours(pms_df: pd.DataFrame, daily_summary: Dict[str, float]) -> pd.DataFrame:
    out = []

    checks = [
        ("ME", "mar", "ME_M03"),
        ("DG1", "mar", "DG1_M03"),
        ("DG2", "mar", "DG2_M03"),
        ("DG3", "mar", "DG3_M03"),
        ("DG4", "mar", "DG4_M03"),
    ]

    for machine, month_col, key in checks:
        expected = daily_summary.get(key)
        if expected is None:
            continue

        if machine == "ME":
            sample = pms_df[pms_df["code_no"].str.startswith("ME-") & pms_df["mar"].notna()]
        else:
            dg_prefix = machine.replace("DG", "DG-0")
            sample = pms_df[pms_df["code_no"].str.startswith(dg_prefix) & pms_df["mar"].notna()]

        if sample.empty:
            out.append({
                "machine": machine,
                "month_col": month_col,
                "expected_from_daily": expected,
                "observed_sample": None,
                "status": "NO_SAMPLE"
            })
            continue

        observed = float(sample.iloc[0][month_col])
        out.append({
            "machine": machine,
            "month_col": month_col,
            "expected_from_daily": expected,
            "observed_sample": observed,
            "status": "OK" if abs(expected - observed) <= 1e-9 else "REVIEW"
        })

    return pd.DataFrame(out)


def strict_quality_gate(
    pms_df: pd.DataFrame,
    ledger_df: pd.DataFrame,
    validation_df: pd.DataFrame
) -> None:
    if pms_df.empty:
        raise SchemaError("PMS extraction produced zero rows.")

    if ledger_df.empty:
        raise ReconciliationError("Ledger build produced zero rows.")

    if pms_df["code_no"].isna().any():
        raise SchemaError("Null code_no values present in PMS output.")

    bad_codes = pms_df[~pms_df["code_no"].astype(str).apply(is_code_like)]
    if not bad_codes.empty:
        raise SchemaError(f"Invalid code_no values leaked into PMS output: {bad_codes['code_no'].head(10).tolist()}")

    if not validation_df.empty and (validation_df["status"] == "REVIEW").any():
        rows = validation_df[validation_df["status"] == "REVIEW"].to_dict(orient="records")
        raise ReconciliationError(f"Daily-hours reconciliation failed: {rows}")


# =========================
# Top-level bundle
# =========================

def build_bundle(
    pms_excel_bytes: bytes,
    pms_excel_name: str,
    running_word_bytes: Optional[bytes] = None,
    running_word_name: Optional[str] = None,
) -> Dict[str, pd.DataFrame]:
    pms_fp = file_fingerprint(pms_excel_name, pms_excel_bytes)

    vessel_daily, daily_df, daily_summary = parse_daily_operating_hours(pms_excel_bytes)
    vessel_pms, pms_df = parse_pms_sheet(pms_excel_bytes)

    if normalize_text(vessel_daily).upper() != normalize_text(vessel_pms).upper():
        raise ReconciliationError(
            f"Vessel mismatch between daily-hours and PMS sheets: '{vessel_daily}' vs '{vessel_pms}'."
        )

    word_df = None
    word_fp = None

    if running_word_bytes is not None and running_word_name is not None:
        word_fp = file_fingerprint(running_word_name, running_word_bytes)
        vessel_word, word_df = parse_word_running_hours(running_word_bytes, running_word_name)

        if normalize_text(vessel_word).upper() != normalize_text(vessel_pms).upper():
            raise ReconciliationError(
                f"Vessel mismatch between PMS and Word report: '{vessel_pms}' vs '{vessel_word}'."
            )

    ledger_df = build_immutable_ledger(pms_df, word_df, pms_fp, word_fp)
    validation_df = validate_pms_vs_daily_hours(pms_df, daily_summary)

    strict_quality_gate(pms_df, ledger_df, validation_df)

    audit_meta = pd.DataFrame([{
        "vessel_name": vessel_pms,
        "pms_file_name": pms_fp.file_name,
        "pms_sha256": pms_fp.sha256,
        "pms_byte_size": pms_fp.byte_size,
        "word_file_name": word_fp.file_name if word_fp else None,
        "word_sha256": word_fp.sha256 if word_fp else None,
        "word_byte_size": word_fp.byte_size if word_fp else None,
        "created_at_utc": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
        "daily_rows": len(daily_df),
        "pms_rows": len(pms_df),
        "word_rows": len(word_df) if word_df is not None else 0,
        "ledger_rows": len(ledger_df),
    }])

    return {
        "daily_hours": daily_df,
        "pms_items": pms_df,
        "word_running_hours": word_df if word_df is not None else pd.DataFrame(),
        "ledger": ledger_df,
        "validation": validation_df,
        "audit_meta": audit_meta,
    }


# =========================
# Export helper
# =========================

def write_bundle_to_excel(bundle: Dict[str, pd.DataFrame], output_path: str) -> None:
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        for sheet_name, df in bundle.items():
            safe_sheet = sheet_name[:31]
            if df is None:
                pd.DataFrame().to_excel(writer, sheet_name=safe_sheet, index=False)
            else:
                df.to_excel(writer, sheet_name=safe_sheet, index=False)


# =========================
# Example usage
# =========================

if __name__ == "__main__":
    PMS_FILE = "TEC-001-PMS-FALCON-Mar.-2026.xlsx"
    WORD_FILE = "TEC-004-RUNNING-HOURS-MONTHLY-REPORT-Mar.-2026.doc"
    OUT_FILE = "immutable_ledger_bundle.xlsx"

    with open(PMS_FILE, "rb") as f:
        pms_bytes = f.read()

    with open(WORD_FILE, "rb") as f:
        word_bytes = f.read()

    bundle = build_bundle(
        pms_excel_bytes=pms_bytes,
        pms_excel_name=PMS_FILE,
        running_word_bytes=word_bytes,
        running_word_name=WORD_FILE,
    )

    write_bundle_to_excel(bundle, OUT_FILE)
    print("OK:", OUT_FILE)
